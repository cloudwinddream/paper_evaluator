"""
论文解析模块
负责读取和解析Word文档，提取文本内容、结构信息等
"""

from __future__ import annotations

import base64
import re
import subprocess
import tempfile
from dataclasses import dataclass, field
from typing import Optional
from pathlib import Path

from docx import Document


def _try_markitdown(filepath: Path) -> str | None:
    """通过 MarkItDown 提取文本（兜底方案，支持 .docx/.doc/.pdf）"""
    try:
        from markitdown import MarkItDown
        md = MarkItDown()
        result = md.convert(str(filepath))
        text = result.text_content
        if text and text.strip():
            return text.strip()
        return None
    except ImportError:
        return None
    except Exception:
        return None


def _extract_text_from_pdf(pdf_path: Path) -> str:
    """从 PDF 文件中提取文本"""
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(pdf_path))
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        return "\n".join(text_parts)
    except Exception as e:
        raise RuntimeError(f"无法解析PDF文件: {pdf_path.name}: {e}")


def _count_images_from_pdf(pdf_path: Path) -> int:
    """从 PDF 文件中提取图片数量（使用 pymupdf）"""
    try:
        import fitz  # pymupdf
        doc = fitz.open(str(pdf_path))
        image_count = 0
        for page in doc:
            image_list = page.get_images(full=True)
            image_count += len(image_list)
        doc.close()
        return image_count
    except ImportError:
        return 0
    except Exception:
        return 0


def _try_via_libreoffice(doc_path: Path) -> Path | None:
    """策略1: 通过 LibreOffice 命令行转换 .doc → .docx（macOS/Linux 首选）"""
    try:
        tmp_dir = Path(tempfile.mkdtemp())
        result = subprocess.run(
            ["soffice", "--headless", "--convert-to", "docx",
             "--outdir", str(tmp_dir), str(doc_path.absolute())],
            capture_output=True, text=True, timeout=60,
        )
        if result.returncode != 0:
            return None
        # soffice 输出文件名与原文件相同（仅扩展名不同）
        out_file = tmp_dir / f"{doc_path.stem}.docx"
        if out_file.exists():
            return out_file
        return None
    except Exception:
        return None


def _count_images_from_doc_via_com(doc_path: Path) -> int:
    """使用 Word COM 直接统计 .doc 文件中的图片数量（避免转换丢失）"""
    try:
        import win32com.client
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(
            str(doc_path.absolute()), ReadOnly=True,
            AddToRecentFiles=False, OpenAndRepair=True,
        )
        count = doc.InlineShapes.Count + doc.Shapes.Count
        doc.Close()
        word.Quit()
        return count
    except Exception:
        try: word.Quit()
        except: pass
        return 0


def _extract_images_from_doc_via_com(doc_path: Path) -> list[ImageData]:
    """使用 Word COM 将 .doc 另存为 HTML，从中提取图片"""
    images: list[ImageData] = []
    try:
        import win32com.client
        import tempfile, shutil
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(
            str(doc_path.absolute()), ReadOnly=True,
            AddToRecentFiles=False, OpenAndRepair=True,
        )
        export_dir = Path(tempfile.mkdtemp())
        html_path = export_dir / "output.html"

        # 另存为筛选过的 HTML（wdFormatFilteredHTML = 10），Word 自动在旁生成图片文件夹
        doc.SaveAs(str(html_path), FileFormat=10)
        doc.Close()
        word.Quit()

        # 收集 HTML 同目录下的图片文件
        img_exts = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".emf", ".wmf"}
        mime_map = {".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
                    ".gif": "image/gif", ".bmp": "image/bmp", ".tiff": "image/tiff",
                    ".emf": "image/x-emf", ".wmf": "image/x-wmf"}
        for f in sorted(export_dir.rglob("*")):
            if not f.is_file() or f.suffix.lower() not in img_exts:
                continue
            with open(f, "rb") as fh:
                img_bytes = fh.read()
            # 过滤过小的图片（Word 导出 HTML 时常产生箭头/分隔线等小图标）
            if len(img_bytes) < 1024:
                continue
            try:
                from PIL import Image as PILImage
                import io
                pil_img = PILImage.open(io.BytesIO(img_bytes))
                w, h = pil_img.size
                if w < 20 or h < 20:
                    continue
            except Exception:
                pass
            b64_str = base64.b64encode(img_bytes).decode("utf-8")
            ext = f.suffix.lower()
            images.append(ImageData(
                index=len(images) + 1,
                section="body",
                caption_context="",
                mime_type=mime_map.get(ext, "image/png"),
                base64_data=b64_str,
            ))

        shutil.rmtree(export_dir, ignore_errors=True)
    except Exception:
        try: word.Quit()
        except: pass
    return images


def _try_via_word(doc_path: Path) -> Path | None:
    """策略2: 通过 Word COM 直接打开（含修复模式）"""
    try:
        import win32com.client
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(
            str(doc_path.absolute()), ReadOnly=True,
            AddToRecentFiles=False, OpenAndRepair=True, Format=0,
        )
        tmp_dir = Path(tempfile.mkdtemp())
        docx_path = tmp_dir / f"{doc_path.stem}.docx"
        doc.SaveAs(str(docx_path.absolute()), FileFormat=16)
        doc.Close()
        word.Quit()
        return docx_path
    except Exception:
        try: word.Quit()
        except: pass
        return None


def _try_via_word_shortpath(doc_path: Path) -> Path | None:
    """策略3: 复制到短路径后通过 Word 打开"""
    try:
        import win32com.client
        import shutil
        short_dir = Path(tempfile.mkdtemp())
        short_path = short_dir / "input.doc"
        shutil.copy2(doc_path, short_path)
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(
            str(short_path.absolute()), ReadOnly=True,
            AddToRecentFiles=False, OpenAndRepair=True, Format=0,
        )
        tmp_dir = Path(tempfile.mkdtemp())
        docx_path = tmp_dir / f"{doc_path.stem}.docx"
        doc.SaveAs(str(docx_path.absolute()), FileFormat=16)
        doc.Close()
        word.Quit()
        shutil.rmtree(short_dir, ignore_errors=True)
        return docx_path
    except Exception:
        try: word.Quit()
        except: pass
        try: shutil.rmtree(short_dir, ignore_errors=True)
        except: pass
        return None


def _try_via_word_as_text(doc_path: Path) -> Path | None:
    """策略4: 通过 Word 以纯文本方式打开（绕过 XML 解析器）"""
    try:
        import win32com.client
        import shutil
        short_dir = Path(tempfile.mkdtemp())
        short_path = short_dir / "input.doc"
        shutil.copy2(doc_path, short_path)
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(
            str(short_path.absolute()), ReadOnly=True,
            AddToRecentFiles=False, Format=7,  # 纯文本
        )
        tmp_dir = Path(tempfile.mkdtemp())
        docx_path = tmp_dir / f"{doc_path.stem}.docx"
        doc.SaveAs(str(docx_path.absolute()), FileFormat=16)
        doc.Close()
        word.Quit()
        shutil.rmtree(short_dir, ignore_errors=True)
        return docx_path
    except Exception:
        try: word.Quit()
        except: pass
        try: shutil.rmtree(short_dir, ignore_errors=True)
        except: pass
        return None


def _try_as_docx(doc_path: Path) -> Path | None:
    """策略5: 尝试直接当 .docx 解析（可能只是扩展名错误）"""
    try:
        import zipfile
        from docx import Document as DocxDocument
        if not zipfile.is_zipfile(doc_path):
            return None
        with zipfile.ZipFile(doc_path) as z:
            if "word/document.xml" not in z.namelist():
                return None
        tmp_dir = Path(tempfile.mkdtemp())
        docx_path = tmp_dir / f"{doc_path.stem}.docx"
        import shutil
        shutil.copy2(doc_path, docx_path)
        DocxDocument(docx_path)
        return docx_path
    except Exception:
        return None


def _try_extract_text_raw(doc_path: Path) -> Path | None:
    """策略6: 从二进制文件中暴力提取可读文本（最后手段）"""
    try:
        text_parts = []
        with open(doc_path, "rb") as f:
            data = f.read()

        # 尝试 UTF-16LE 解码（Word 内部编码）
        i = 0
        current = []
        while i < len(data) - 1:
            try:
                char = data[i:i+2].decode("utf-16-le")
                if char.isprintable() or char in "\n\r\t":
                    current.append(char)
                else:
                    if len("".join(current)) > 20:
                        text_parts.append("".join(current))
                    current = []
            except:
                if len("".join(current)) > 20:
                    text_parts.append("".join(current))
                current = []
                i += 1
                continue
            i += 2

        text = "\n".join(text_parts)
        if len(text) < 50:
            return None

        tmp_dir = Path(tempfile.mkdtemp())
        docx_path = tmp_dir / f"{doc_path.stem}.docx"

        from docx import Document
        doc = Document()
        for line in text.split("\n"):
            line = line.strip()
            if line:
                doc.add_paragraph(line)
        doc.save(docx_path)
        return docx_path
    except Exception:
        return None


def _convert_doc_to_docx(doc_path: Path) -> Path:
    """将 .doc 文件转换为临时的 .docx 文件（自动尝试多种修复策略）"""
    strategies = [
        ("LibreOffice 命令行转换", _try_via_libreoffice),
        ("Word 修复模式打开", _try_via_word),
        ("复制到短路径后重试", _try_via_word_shortpath),
        ("纯文本方式打开", _try_via_word_as_text),
        ("尝试作为 .docx 解析", _try_as_docx),
        ("暴力提取文本", _try_extract_text_raw),
    ]
    for name, func in strategies:
        result = func(doc_path)
        if result is not None:
            return result

    raise RuntimeError(f"无法解析文件: {doc_path.name}（所有修复策略均失败）")


@dataclass
class ImageData:
    """提取的图片数据结构"""
    index: int                             # 图片编号（从1开始）
    section: str                           # 所在章节名称
    caption_context: str                   # 上下文文本（~200字符）
    mime_type: str                         # MIME 类型 (如 "image/png")
    base64_data: str                       # base64 编码的图片数据


@dataclass
class ParsedPaper:
    """解析后的论文数据结构"""
    filename: str                          # 文件名
    student_name: str = ""                 # 学生姓名（从文件名提取）
    student_id: str = ""                   # 学号（如果有）
    title: str = ""                        # 论文标题
    raw_text: str = ""                     # 原始文本
    sections: dict = field(default_factory=dict)  # 各部分内容
    word_count: int = 0                    # 字数
    paragraph_count: int = 0               # 段落数
    figure_count: int = 0                  # 图片数量
    table_count: int = 0                   # 表格数量
    reference_count: int = 0               # 参考文献数量
    has_abstract: bool = False             # 是否有摘要
    has_keywords: bool = False             # 是否有关键词
    has_references: bool = False           # 是否有参考文献
    images: list[ImageData] = field(default_factory=list)  # 提取的图片列表


class PaperParser:
    """Word论文解析器"""

    # 常见章节标题模式
    SECTION_PATTERNS = {
        "abstract": [r"摘\s*要", r"Abstract"],
        "keywords": [r"关\s*键\s*词", r"Keywords"],
        "introduction": [r"引\s*言", r"绪\s*论", r"前\s*言", r"第\s*1\s*章", r"Chapter\s*1"],
        "body": [r"正\s*文", r"第\s*[2-9]\s*章", r"Chapter\s*[2-9]"],
        "conclusion": [r"结\s*论", r"总\s*结", r"结\s*语", r"Conclusion"],
        "references": [r"参\s*考\s*文\s*献", r"References", r"Bibliography"],
    }

    def __init__(self):
        self.papers: list[ParsedPaper] = []

    def parse_file(self, filepath: str | Path) -> ParsedPaper:
        """解析单个文件（支持 .docx, .doc 和 .pdf）"""
        filepath = Path(filepath)
        suffix = filepath.suffix.lower()

        if suffix == ".pdf":
            return self._parse_pdf(filepath)

        is_doc = suffix == ".doc"
        temp_docx = None
        source_path = filepath

        if is_doc:
            temp_docx = _convert_doc_to_docx(filepath)
            source_path = temp_docx

        try:
            try:
                doc = Document(source_path)

                paper = ParsedPaper(
                    filename=filepath.name,
                    student_name=self._extract_student_name(filepath.name),
                )

                # 提取纯文本
                paper.raw_text = self._extract_text(doc)

                # 统计字数（中文字符+英文单词）
                paper.word_count = self._count_words(paper.raw_text)

                # 段落数
                paper.paragraph_count = len(doc.paragraphs)

                # 图片数量 & 提取（.doc 用 COM 直接读取，.docx 用 python-docx）
                if is_doc:
                    paper.figure_count = _count_images_from_doc_via_com(filepath)
                    paper.images = _extract_images_from_doc_via_com(filepath)
                    if not paper.images and paper.figure_count > 0:
                        # COM 提取失败，改用转换后的 docx
                        paper.images = self._extract_images_from_docx(doc)
                        paper.figure_count = max(paper.figure_count, self._count_figures(doc))
                else:
                    paper.figure_count = self._count_figures(doc)
                    paper.images = self._extract_images_from_docx(doc)

                # 表格数量
                paper.table_count = len(doc.tables)

                # 提取各部分内容
                paper.sections = self._extract_sections(doc)

                # 检查必要部分
                paper.has_abstract = bool(paper.sections.get("abstract"))
                paper.has_keywords = bool(paper.sections.get("keywords"))
                paper.has_references = bool(paper.sections.get("references"))

                # 统计参考文献数量
                paper.reference_count = self._count_references(paper.sections.get("references", ""))

                # 提取标题（通常是文档第一个有意义的段落）
                paper.title = self._extract_title(doc)

                return paper
            except Exception:
                raw_text = _try_markitdown(filepath)
                if raw_text:
                    return self._build_paper_from_text(filepath, raw_text)
                raise
        finally:
            if temp_docx and temp_docx.exists():
                temp_docx.unlink()
                try:
                    temp_docx.parent.rmdir()
                except OSError:
                    pass

    def _build_paper_from_text(
        self, filepath: Path, raw_text: str, figure_count: int = 0,
        images: list[ImageData] | None = None,
    ) -> ParsedPaper:
        """从纯文本构建 ParsedPaper（兜底降级用）"""
        lines = [l for l in raw_text.split("\n") if l.strip()]
        paper = ParsedPaper(
            filename=filepath.name,
            student_name=self._extract_student_name(filepath.name),
            raw_text=raw_text,
            word_count=self._count_words(raw_text),
            paragraph_count=len(lines),
            figure_count=figure_count,
        )
        if images is not None:
            paper.images = images
        paper.sections = self._extract_sections_from_text(raw_text)
        paper.has_abstract = bool(paper.sections.get("abstract"))
        paper.has_keywords = bool(paper.sections.get("keywords"))
        paper.has_references = bool(paper.sections.get("references"))
        paper.reference_count = self._count_references(paper.sections.get("references", ""))
        paper.title = self._extract_title_from_text(raw_text, lines)
        return paper

    def _parse_pdf(self, filepath: Path) -> ParsedPaper:
        """解析 PDF 文件（pypdf + pymupdf 图片计数，失败后降级到 MarkItDown）"""
        raw_text = None
        try:
            raw_text = _extract_text_from_pdf(filepath)
        except Exception:
            raw_text = _try_markitdown(filepath)
        if not raw_text:
            raise RuntimeError(f"无法解析PDF文件: {filepath.name}（pypdf + MarkItDown 均失败）")

        # 尝试用 pymupdf 统计图片数量并提取图片数据
        figure_count = _count_images_from_pdf(filepath)
        images = self._extract_images_from_pdf(filepath)
        return self._build_paper_from_text(filepath, raw_text, figure_count, images=images)

    def _extract_and_rename_zips(self, directory: Path) -> list[Path]:
        """解压目录下的所有 .zip 文件，并将内部文件重命名为压缩包的名字"""
        extracted: list[Path] = []
        for zip_path in sorted(directory.glob("*.zip")):
            stem = zip_path.stem
            import zipfile
            try:
                with zipfile.ZipFile(zip_path, "r") as zf:
                    # 解压到同名临时子目录，避免覆盖已有文件
                    extract_dir = directory / f".{stem}_extracted"
                    extract_dir.mkdir(parents=True, exist_ok=True)
                    zf.extractall(extract_dir)

                    for f in sorted(extract_dir.rglob("*")):
                        if not f.is_file():
                            continue
                        # 跳过常见的非文档文件
                        if f.suffix.lower() not in (".docx", ".doc", ".pdf"):
                            continue
                        new_name = stem + f.suffix
                        dest = directory / new_name
                        # 如果目标已存在则跳过
                        if dest.exists():
                            print(f"[解压提示] {new_name} 已存在，跳过")
                            continue
                        f.rename(dest)
                        extracted.append(dest)
                        print(f"[解压成功] {zip_path.name} → {new_name}")

                    # 清理临时目录
                    import shutil
                    shutil.rmtree(extract_dir, ignore_errors=True)
            except Exception as e:
                print(f"[解压失败] {zip_path.name}: {e}")

            # 关闭文件后再删除压缩包
            if zip_path.exists():
                try:
                    zip_path.unlink()
                    print(f"[解压完成] 已删除压缩包: {zip_path.name}")
                except Exception as e:
                    print(f"[解压提示] 删除压缩包失败: {zip_path.name}: {e}")
        return extracted

    def parse_directory(self, directory: str | Path) -> list[ParsedPaper]:
        """解析目录下所有支持的文件（.docx, .doc, .pdf）"""
        directory = Path(directory)
        # 先处理 zip 压缩包
        self._extract_and_rename_zips(directory)
        self.papers = []

        docx_names = {f.stem for f in directory.glob("*.docx")}
        for filepath in sorted(directory.glob("*.docx")) + sorted(directory.glob("*.doc")) + sorted(directory.glob("*.pdf")):
            # 同名的 .doc 和 .docx 同时存在时，跳过 .doc 避免重复
            if filepath.suffix.lower() == ".doc" and filepath.stem in docx_names:
                continue
            try:
                paper = self.parse_file(filepath)
                self.papers.append(paper)
                print(f"[解析成功] {filepath.name} - {paper.word_count}字")
            except Exception as e:
                print(f"[解析失败] {filepath.name}: {e}")

        return self.papers

    def _extract_text(self, doc: Document) -> str:
        """提取文档纯文本"""
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        return "\n".join(paragraphs)

    def _count_words(self, text: str) -> int:
        """统计字数（中文字符数 + 英文单词数）"""
        # 中文字符数
        chinese_chars = len(re.findall(r"[一-鿿]", text))
        # 英文单词数
        english_words = len(re.findall(r"[a-zA-Z]+", text))
        return chinese_chars + english_words

    def _count_figures(self, doc: Document) -> int:
        """统计图片数量"""
        image_count = 0
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                image_count += 1
        return image_count

    def _extract_images_from_docx(self, doc: Document) -> list[ImageData]:
        """从 docx Document 对象中提取图片

        遍历每个段落，检查 XML 中是否有 a:blip 元素（表示嵌入的图片）。
        追踪当前章节，提取上下文文本作为 caption_context。
        """
        images: list[ImageData] = []
        current_section = "header"

        try:
            ns_a = "http://schemas.openxmlformats.org/drawingml/2006/main"
            ns_r = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

            paragraphs = doc.paragraphs
            for i, para in enumerate(paragraphs):
                text = para.text.strip()

                # 追踪当前章节（复用章节标题检测逻辑）
                if text:
                    matched_section = None
                    for section_name, patterns in self.SECTION_PATTERNS.items():
                        for pattern in patterns:
                            if re.match(pattern, text, re.IGNORECASE):
                                matched_section = section_name
                                break
                        if matched_section:
                            break
                    if matched_section:
                        current_section = matched_section

                # 在段落 XML 中查找 a:blip 元素（图片引用）
                blips = para._element.findall(f".//{{{ns_a}}}blip")
                if not blips:
                    continue

                for blip in blips:
                    # 获取关系 ID (r:embed 属性)
                    r_embed = blip.get(f"{{{ns_r}}}embed")
                    if not r_embed:
                        continue

                    # 获取图片二进制数据
                    try:
                        image_part = doc.part.related_parts[r_embed]
                        image_bytes = image_part.blob
                    except (KeyError, AttributeError):
                        continue

                    # 确定 MIME 类型
                    content_type = getattr(image_part, "content_type", "")
                    if not content_type:
                        # 根据关系 ID 后缀推断
                        if r_embed.lower().endswith((".png", ".PNG")):
                            content_type = "image/png"
                        elif r_embed.lower().endswith((".jpg", ".jpeg", ".JPG", ".JPEG")):
                            content_type = "image/jpeg"
                        else:
                            content_type = "image/png"

                    # 提取 caption_context
                    if text:
                        caption_context = text[:200]
                    elif i + 1 < len(paragraphs):
                        next_text = paragraphs[i + 1].text.strip()
                        if next_text and len(next_text) < 200:
                            caption_context = next_text[:200]
                        else:
                            caption_context = current_section
                    else:
                        caption_context = current_section

                    # 转为 base64
                    b64_str = base64.b64encode(image_bytes).decode("utf-8")

                    images.append(ImageData(
                        index=len(images) + 1,
                        section=current_section,
                        caption_context=caption_context,
                        mime_type=content_type,
                        base64_data=b64_str,
                    ))

        except Exception as e:
            print(f"[图片提取警告] docx 图片提取失败: {e}")
            return []

        return images

    def _extract_images_from_pdf(self, pdf_path: Path) -> list[ImageData]:
        """从 PDF 文件中提取图片（使用 pymupdf/fitz）

        遍历每一页，使用 page.get_images() 获取图片引用，
        通过 doc.extract_image(xref) 获取图片二进制数据。
        建立页面→章节映射以确定图片所在章节。
        最多提取 10 张图片。
        """
        import fitz  # pymupdf
        images: list[ImageData] = []

        try:
            doc = fitz.open(str(pdf_path))

            # 构建页面→章节映射
            page_sections: dict[int, str] = {}
            current_section = "header"
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()
                for line in text.split("\n"):
                    line = line.strip()
                    if not line:
                        continue
                    matched_section = None
                    for section_name, patterns in self.SECTION_PATTERNS.items():
                        for pattern in patterns:
                            if re.match(pattern, line, re.IGNORECASE):
                                matched_section = section_name
                                break
                        if matched_section:
                            break
                    if matched_section:
                        current_section = matched_section
                page_sections[page_num] = current_section

            # 扩展名 → MIME 类型映射
            mime_map = {
                "png": "image/png",
                "jpg": "image/jpeg",
                "jpeg": "image/jpeg",
                "gif": "image/gif",
                "bmp": "image/bmp",
                "tiff": "image/tiff",
                "tif": "image/tiff",
                "webp": "image/webp",
            }

            for page_num in range(len(doc)):
                page = doc[page_num]
                image_list = page.get_images(full=True)

                for img_ref in image_list:
                    xref = img_ref[0]
                    try:
                        img_info = doc.extract_image(xref)
                    except Exception:
                        continue

                    img_bytes = img_info["image"]
                    img_ext = img_info.get("ext", "png").lower()
                    mime_type = mime_map.get(img_ext, f"image/{img_ext}")

                    b64_str = base64.b64encode(img_bytes).decode("utf-8")

                    # 提取上下文（页面第一段有意义的文本）
                    page_text_lines = [
                        l.strip() for l in page.get_text().split("\n")
                        if l.strip()
                    ]
                    caption_context = (
                        page_text_lines[0][:200] if page_text_lines else current_section
                    )

                    images.append(ImageData(
                        index=len(images) + 1,
                        section=page_sections.get(page_num, "header"),
                        caption_context=caption_context,
                        mime_type=mime_type,
                        base64_data=b64_str,
                    ))

            doc.close()
        except ImportError:
            print("[图片提取警告] pymupdf 未安装，无法提取 PDF 图片")
            return []
        except Exception as e:
            print(f"[图片提取警告] PDF 图片提取失败: {e}")
            return []

        return images

    def _extract_sections(self, doc: Document) -> dict[str, str]:
        """提取各章节内容"""
        sections = {}
        current_section = "header"
        current_content = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # 检查是否是章节标题
            matched_section = None
            for section_name, patterns in self.SECTION_PATTERNS.items():
                for pattern in patterns:
                    if re.match(pattern, text, re.IGNORECASE):
                        matched_section = section_name
                        break
                if matched_section:
                    break

            if matched_section:
                # 保存前一个章节的内容
                if current_content:
                    sections[current_section] = "\n".join(current_content)
                current_section = matched_section
                current_content = []
            else:
                current_content.append(text)

        # 保存最后一个章节
        if current_content:
            sections[current_section] = "\n".join(current_content)

        return sections

    def _extract_sections_from_text(self, text: str) -> dict[str, str]:
        """从纯文本中提取各章节内容"""
        sections = {}
        current_section = "header"
        current_content = []

        for line in text.split("\n"):
            line = line.strip()
            if not line:
                continue

            matched_section = None
            for section_name, patterns in self.SECTION_PATTERNS.items():
                for pattern in patterns:
                    if re.match(pattern, line, re.IGNORECASE):
                        matched_section = section_name
                        break
                if matched_section:
                    break

            if matched_section:
                if current_content:
                    sections[current_section] = "\n".join(current_content)
                current_section = matched_section
                current_content = []
            else:
                current_content.append(line)

        if current_content:
            sections[current_section] = "\n".join(current_content)

        return sections

    def _count_references(self, references_text: str) -> int:
        """统计参考文献数量"""
        if not references_text:
            return 0
        # 通常参考文献以数字编号，如 [1] [2] ...
        matches = re.findall(r"\[\d+\]", references_text)
        if matches:
            return len(matches)
        # 如果没有编号，按行数估算
        lines = [l for l in references_text.split("\n") if l.strip()]
        return len(lines)

    def _extract_student_name(self, filename: str) -> str:
        """从文件名提取学生标识
        格式1: 学部-专业-班级-学号-姓名  → 学号-姓名
        格式2: 学号_姓名                → 学号-姓名
        格式3: 姓名_题目                → 姓名
        格式4: 姓名                    → 姓名
        """
        name = Path(filename).stem
        parts = [p for p in re.split(r"[_\-\s]+", name) if p]

        # 找出所有数字段（学号）和非数字段
        digits = [(i, p) for i, p in enumerate(parts) if p.isdigit() and len(p) >= 6]
        names = [(i, p) for i, p in enumerate(parts) if not p.isdigit() and len(p) >= 2]

        if digits:
            # 有学号 → 取最后一个学号及它之后的第一个姓名
            sid_idx, sid = digits[-1]
            sname = parts[-1]  # 默认取最后一段
            for i, p in names:
                if i > sid_idx:
                    sname = p
                    break
            return f"{sid}-{sname}"

        # 无学号 → 取第一个非数字段作为姓名
        if names:
            return names[0][1]

        return name

    def _extract_title(self, doc: Document) -> str:
        """提取论文标题"""
        for para in doc.paragraphs:
            text = para.text.strip()
            if text and len(text) < 100:
                # 跳过"论文题目"等标签
                if re.match(r"论文题目|题目|Title", text, re.IGNORECASE):
                    continue
                return text
        return "未识别标题"

    def _extract_title_from_text(self, text: str, lines: list[str] | None = None) -> str:
        """从纯文本中提取标题"""
        if lines is None:
            lines = [l.strip() for l in text.split("\n") if l.strip()]
        for line in lines:
            if len(line) < 100:
                if re.match(r"论文题目|题目|Title", line, re.IGNORECASE):
                    continue
                return line
        return "未识别标题"

    @staticmethod
    def parse_requirements_doc(filepath: str | Path) -> str:
        """读取题目要求文档（支持 .docx, .doc, .pdf），返回纯文本"""
        filepath = Path(filepath)
        if not filepath.exists():
            raise FileNotFoundError(f"题目要求文档不存在: {filepath}")

        suffix = filepath.suffix.lower()
        if suffix == ".pdf":
            try:
                raw_text = _extract_text_from_pdf(filepath)
            except Exception:
                raw_text = _try_markitdown(filepath)
            if not raw_text:
                raise RuntimeError(f"无法解析题目要求文档: {filepath.name}")
            lines = [l.strip() for l in raw_text.split("\n") if l.strip()]
            return "\n".join(lines)

        is_doc = suffix == ".doc"
        temp_docx = None
        source_path = filepath

        if is_doc:
            temp_docx = _convert_doc_to_docx(filepath)
            source_path = temp_docx

        try:
            try:
                doc = Document(source_path)
                paragraphs = []
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if text:
                        paragraphs.append(text)
                return "\n".join(paragraphs)
            except Exception:
                raw_text = _try_markitdown(filepath)
                if raw_text:
                    return raw_text
                raise
        finally:
            if temp_docx and temp_docx.exists():
                temp_docx.unlink()
                try:
                    temp_docx.parent.rmdir()
                except OSError:
                    pass
