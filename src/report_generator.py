"""
报告生成模块
生成Excel汇总表、Word详细报告和Markdown报告
"""

from datetime import datetime
from pathlib import Path

from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.paper_parser import ParsedPaper
from src.completeness_checker import CompletenessResult
from src.ai_evaluator import EvaluationResult
from src.aigc_detector import AIGCResult
from src.plagiarism_checker import PlagiarismResult


class ReportGenerator:
    """报告生成器"""

    def __init__(self, output_dir: str):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    # ──────────────────────────────────────────────
    # Excel 汇总表
    # ──────────────────────────────────────────────

    def generate_excel(
        self,
        papers: list[ParsedPaper],
        completeness_results: list[CompletenessResult],
        evaluation_results: list[EvaluationResult],
        aigc_results: list[AIGCResult],
        plagiarism_results: list[PlagiarismResult],
        dimensions: list[dict],
        filename: str = "scores_summary.xlsx",
    ) -> str:
        """生成Excel汇总表（每个维度单独一列，分数取整）"""
        wb = Workbook()
        ws = wb.active
        ws.title = "评分汇总"

        # ── 样式 ──
        header_font = Font(name="微软雅黑", bold=True, size=11, color="FFFFFF")
        header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        header_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
        data_font = Font(name="微软雅黑", size=10)
        data_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
        left_align = Alignment(horizontal="left", vertical="center", wrap_text=True)
        thin_border = Border(
            left=Side(style="thin"), right=Side(style="thin"),
            top=Side(style="thin"), bottom=Side(style="thin"),
        )

        # ── 动态列：基础列 + 每个评分维度 + 综合列 ──
        base_headers = ["序号", "学生姓名", "论文标题", "字数", "完整性得分"]
        dim_headers = [d["name"] for d in dimensions]
        tail_headers = ["AI评审总分", "AIGC风险", "查重最高相似度", "最终得分", "简短评语", "查重警告"]
        all_headers = base_headers + dim_headers + tail_headers

        for col, header in enumerate(all_headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_align
            cell.border = thin_border

        # ── 结果映射 ──
        completeness_map = {r.student_name: r for r in completeness_results}
        evaluation_map = {r.student_name: r for r in evaluation_results}
        aigc_map = {r.student_name: r for r in aigc_results}
        plagiarism_map = {r.student_name: r for r in plagiarism_results}

        # ── 数据行 ──
        for idx, paper in enumerate(papers, 1):
            row = idx + 1
            comp = completeness_map.get(paper.student_name)
            eval_ = evaluation_map.get(paper.student_name)
            aigc = aigc_map.get(paper.student_name)
            plag = plagiarism_map.get(paper.student_name)

            # 最终得分（强制 60-89）
            final_score = 0
            if comp and eval_ and eval_.success:
                final_score = comp.score * 0.2 + eval_.total_score * 0.8
                if aigc and aigc.is_suspicious:
                    final_score *= (1 - aigc.ai_probability * 0.3)
                if plag and plag.highest_similarity > 0.3:
                    final_score *= (1 - plag.highest_similarity * 0.2)
                final_score = min(89, max(60, round(final_score)))
            elif comp and not eval_:
                final_score = min(89, max(60, round(comp.score)))

            # 查重警告
            plag_warning = ""
            if plag and plag.suspicious_pairs:
                similar = [p.student_b if p.student_a == paper.student_name
                           else p.student_a for p in plag.suspicious_pairs]
                plag_warning = f"与{', '.join(similar)}相似度{plag.highest_similarity:.0%}"

            # 组装行数据
            base_data = [
                idx,
                paper.student_name,
                paper.title,
                paper.word_count,
                round(comp.score) if comp else "N/A",
            ]
            # 每个维度的分数
            dim_data = []
            for d in dimensions:
                score = eval_.dimension_scores.get(d["name"], "N/A") if eval_ and eval_.success else "N/A"
                dim_data.append(score)
            tail_data = [
                eval_.total_score if eval_ and eval_.success else "评审失败",
                aigc.overall_risk if aigc else "N/A",
                f"{plag.highest_similarity:.0%}" if plag else "0%",
                final_score,
                eval_.short_comment if eval_ else "",
                plag_warning,
            ]
            row_data = base_data + dim_data + tail_data

            for col, value in enumerate(row_data, 1):
                cell = ws.cell(row=row, column=col, value=value)
                cell.font = data_font
                cell.border = thin_border
                # 姓名、标题、评语、警告左对齐
                if col in [2, 3, len(all_headers) - 1, len(all_headers)]:
                    cell.alignment = left_align
                else:
                    cell.alignment = data_align

                # 最终得分颜色标记
                if col == len(all_headers) - 2 and isinstance(value, int):
                    if value >= 90:
                        cell.font = Font(name="微软雅黑", size=10, color="00B050", bold=True)
                    elif value >= 60:
                        cell.font = Font(name="微软雅黑", size=10, color="000000")
                    else:
                        cell.font = Font(name="微软雅黑", size=10, color="FF0000", bold=True)

        # ── 列宽 ──
        col_widths = [6, 12, 30, 8, 14]  # 基础列
        col_widths += [14] * len(dimensions)  # 每个维度
        col_widths += [12, 10, 14, 10, 30, 25]  # 尾部列
        for i, w in enumerate(col_widths, 1):
            ws.column_dimensions[get_column_letter(i)].width = w

        # ── 统计信息 ──
        stats_row = len(papers) + 3
        ws.cell(row=stats_row, column=1, value="统计信息").font = Font(name="微软雅黑", bold=True, size=11)
        if evaluation_results:
            scores = [r.total_score for r in evaluation_results if r.success]
            if scores:
                ws.cell(row=stats_row + 1, column=1, value=f"平均分: {sum(scores) / len(scores):.0f}")
                ws.cell(row=stats_row + 2, column=1, value=f"最高分: {max(scores)}")
                ws.cell(row=stats_row + 3, column=1, value=f"最低分: {min(scores)}")

        # ── 维度说明 sheet ──
        ws2 = wb.create_sheet("评分维度说明")
        ws2.cell(row=1, column=1, value="维度名称").font = header_font
        ws2.cell(row=1, column=1).fill = header_fill
        ws2.cell(row=1, column=2, value="权重").font = header_font
        ws2.cell(row=1, column=2).fill = header_fill
        ws2.cell(row=1, column=3, value="说明").font = header_font
        ws2.cell(row=1, column=3).fill = header_fill
        for i, d in enumerate(dimensions, 2):
            ws2.cell(row=i, column=1, value=d["name"])
            ws2.cell(row=i, column=2, value=f"{d['weight']*100:.0f}%")
            ws2.cell(row=i, column=3, value=d["description"])
        ws2.column_dimensions["A"].width = 18
        ws2.column_dimensions["B"].width = 10
        ws2.column_dimensions["C"].width = 50

        output_path = self.output_dir / filename
        wb.save(output_path)
        print(f"[Excel报告] 已保存至: {output_path}")
        return str(output_path)

    # ──────────────────────────────────────────────
    # Word 详细报告
    # ──────────────────────────────────────────────

    def generate_word_report(
        self,
        papers: list[ParsedPaper],
        completeness_results: list[CompletenessResult],
        evaluation_results: list[EvaluationResult],
        aigc_results: list[AIGCResult],
        plagiarism_results: list[PlagiarismResult],
        requirements: str,
        filename: str = "evaluation_report.docx",
    ) -> str:
        """生成Word详细报告"""
        doc = Document()

        style = doc.styles["Normal"]
        font = style.font
        font.name = "微软雅黑"
        font.size = Pt(10.5)

        title = doc.add_heading("论文评审报告", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_heading("基本信息", level=1)
        doc.add_paragraph(f"评审时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph(f"论文数量: {len(papers)}篇")

        doc.add_heading("题目要求", level=1)
        doc.add_paragraph(requirements)

        completeness_map = {r.student_name: r for r in completeness_results}
        evaluation_map = {r.student_name: r for r in evaluation_results}
        aigc_map = {r.student_name: r for r in aigc_results}
        plagiarism_map = {r.student_name: r for r in plagiarism_results}

        doc.add_heading("详细评审结果", level=1)

        for paper in papers:
            comp = completeness_map.get(paper.student_name)
            eval_ = evaluation_map.get(paper.student_name)
            aigc = aigc_map.get(paper.student_name)
            plag = plagiarism_map.get(paper.student_name)

            doc.add_heading(f"{paper.student_name} - {paper.title}", level=2)

            table = doc.add_table(rows=4, cols=2, style="Light Grid Accent 1")
            for i, (k, v) in enumerate([
                ("文件名", paper.filename),
                ("字数", f"{paper.word_count}字"),
                ("段落数", str(paper.paragraph_count)),
                ("图片/表格", f"{paper.figure_count}/{paper.table_count}"),
            ]):
                table.rows[i].cells[0].text = k
                table.rows[i].cells[1].text = v
            doc.add_paragraph("")

            doc.add_heading("完整性检测", level=3)
            if comp:
                doc.add_paragraph(f"完整性得分: {comp.score:.0f}/100")
                if comp.missing_sections:
                    doc.add_paragraph(f"缺失部分: {', '.join(comp.missing_sections)}")
                for w in comp.warnings:
                    doc.add_paragraph(f"⚠ {w}", style="List Bullet")

            doc.add_heading("AI评审", level=3)
            if eval_ and eval_.success:
                doc.add_paragraph(f"总分: {eval_.total_score}/100")
                if eval_.dimension_scores:
                    t = doc.add_table(rows=len(eval_.dimension_scores) + 1, cols=2, style="Light Grid Accent 1")
                    t.rows[0].cells[0].text = "维度"
                    t.rows[0].cells[1].text = "得分"
                    for i, (dim, score) in enumerate(eval_.dimension_scores.items(), 1):
                        t.rows[i].cells[0].text = dim
                        t.rows[i].cells[1].text = str(score)
                doc.add_paragraph("")
                doc.add_paragraph(f"评语: {eval_.short_comment}")
                doc.add_paragraph("评分依据:")
                doc.add_paragraph(eval_.evaluation_basis)
            else:
                doc.add_paragraph(f"评审失败: {eval_.error_message if eval_ else '未知错误'}")

            doc.add_heading("AIGC检测", level=3)
            if aigc:
                doc.add_paragraph(f"AI生成概率: {aigc.ai_probability:.0%}")
                doc.add_paragraph(f"风险等级: {aigc.overall_risk}")
                for seg in aigc.suspicious_segments[:5]:
                    doc.add_paragraph(f"• [{seg['type']}] {seg['context']}", style="List Bullet")

            doc.add_heading("查重结果", level=3)
            if plag and plag.suspicious_pairs:
                doc.add_paragraph(f"最高相似度: {plag.highest_similarity:.0%}")
                doc.add_paragraph(f"最相似学生: {plag.most_similar_student}")
                for pair in plag.suspicious_pairs:
                    other = pair.student_b if pair.student_a == paper.student_name else pair.student_a
                    doc.add_paragraph(f"与 {other} 相似度: {pair.similarity:.0%}", style="List Bullet")
            else:
                doc.add_paragraph("未发现明显重复")

            doc.add_page_break()

        doc.add_heading("查重总结", level=1)
        cases = [r for r in plagiarism_results if r.suspicious_pairs]
        if cases:
            doc.add_paragraph(f"共发现 {len(cases)} 名学生存在疑似抄袭情况:")
            for c in cases:
                for pair in c.suspicious_pairs:
                    other = pair.student_b if pair.student_a == c.student_name else pair.student_a
                    doc.add_paragraph(f"• {c.student_name} 与 {other} 相似度 {pair.similarity:.0%}", style="List Bullet")
        else:
            doc.add_paragraph("未发现明显抄袭情况")

        output_path = self.output_dir / filename
        doc.save(output_path)
        print(f"[Word报告] 已保存至: {output_path}")
        return str(output_path)

    # ──────────────────────────────────────────────
    # Markdown 报告（不含题目要求）
    # ──────────────────────────────────────────────

    def generate_markdown_report(
        self,
        papers: list[ParsedPaper],
        completeness_results: list[CompletenessResult],
        evaluation_results: list[EvaluationResult],
        aigc_results: list[AIGCResult],
        plagiarism_results: list[PlagiarismResult],
        dimensions: list[dict],
        filename: str = "evaluation_report.md",
    ) -> str:
        """生成Markdown评审报告（不含题目要求）"""
        completeness_map = {r.student_name: r for r in completeness_results}
        evaluation_map = {r.student_name: r for r in evaluation_results}
        aigc_map = {r.student_name: r for r in aigc_results}
        plagiarism_map = {r.student_name: r for r in plagiarism_results}
        dim_names = [d["name"] for d in dimensions]

        lines: list[str] = []
        now = datetime.now().strftime("%Y-%m-%d %H:%M")

        # ── 标题 ──
        lines.append("# 论文评审报告")
        lines.append("")
        lines.append(f"> 评审时间：{now}　｜　论文数量：{len(papers)}篇")
        lines.append("")

        # ── 评分维度说明 ──
        lines.append("## 评分维度")
        lines.append("")
        lines.append("| 维度 | 权重 | 说明 |")
        lines.append("|------|------|------|")
        for d in dimensions:
            lines.append(f"| {d['name']} | {d['weight']*100:.0f}% | {d['description']} |")
        lines.append("")

        # ── 汇总表 ──
        lines.append("## 评分汇总")
        lines.append("")
        header_cols = ["序号", "学生姓名", "字数", "完整性"] + dim_names + ["总分", "AIGC风险", "查重", "最终得分", "评语"]
        lines.append("| " + " | ".join(header_cols) + " |")
        lines.append("| " + " | ".join(["---"] * len(header_cols)) + " |")

        for idx, paper in enumerate(papers, 1):
            comp = completeness_map.get(paper.student_name)
            eval_ = evaluation_map.get(paper.student_name)
            aigc = aigc_map.get(paper.student_name)
            plag = plagiarism_map.get(paper.student_name)

            final_score = 0
            if comp and eval_ and eval_.success:
                final_score = comp.score * 0.2 + eval_.total_score * 0.8
                if aigc and aigc.is_suspicious:
                    final_score *= (1 - aigc.ai_probability * 0.3)
                if plag and plag.highest_similarity > 0.3:
                    final_score *= (1 - plag.highest_similarity * 0.2)
                final_score = min(89, max(60, round(final_score)))
            elif comp and not eval_:
                final_score = min(89, max(60, round(comp.score)))

            dim_scores = [str(eval_.dimension_scores.get(dn, "-")) if eval_ and eval_.success else "-" for dn in dim_names]
            row = [
                str(idx),
                paper.student_name,
                str(paper.word_count),
                str(round(comp.score)) if comp else "-",
            ] + dim_scores + [
                str(eval_.total_score) if eval_ and eval_.success else "评审失败",
                aigc.overall_risk if aigc else "-",
                f"{plag.highest_similarity:.0%}" if plag else "0%",
                str(final_score),
                eval_.short_comment if eval_ else "-",
            ]
            lines.append("| " + " | ".join(row) + " |")

        lines.append("")

        # ── 统计 ──
        eval_scores = [r.total_score for r in evaluation_results if r.success]
        if eval_scores:
            lines.append("### 统计信息")
            lines.append("")
            lines.append(f"- 平均分：**{sum(eval_scores) / len(eval_scores):.0f}**")
            lines.append(f"- 最高分：**{max(eval_scores)}**")
            lines.append(f"- 最低分：**{min(eval_scores)}**")
            lines.append("")

        # ── 查重总结 ──
        lines.append("## 查重总结")
        lines.append("")
        plag_cases = [r for r in plagiarism_results if r.suspicious_pairs]
        if plag_cases:
            lines.append(f"⚠ 共发现 **{len(plag_cases)}** 名学生存在疑似抄袭情况：")
            lines.append("")
            for c in plag_cases:
                for pair in c.suspicious_pairs:
                    other = pair.student_b if pair.student_a == c.student_name else pair.student_a
                    lines.append(f"- **{c.student_name}** 与 **{other}** 相似度 {pair.similarity:.0%}")
        else:
            lines.append("✅ 未发现明显抄袭情况")
        lines.append("")

        # ── AIGC 风险 ──
        lines.append("## AIGC 检测总结")
        lines.append("")
        suspicious_aigc = [r for r in aigc_results if r.is_suspicious]
        if suspicious_aigc:
            lines.append(f"⚠ 共发现 **{len(suspicious_aigc)}** 篇论文存在AI生成风险：")
            lines.append("")
            for r in suspicious_aigc:
                lines.append(f"- **{r.student_name}**：{r.overall_risk}（AI概率 {r.ai_probability:.0%}）")
        else:
            lines.append("✅ 未发现明显AI生成内容")
        lines.append("")

        # ── 逐个学生详情 ──
        lines.append("---")
        lines.append("")
        lines.append("## 逐个学生评审详情")
        lines.append("")

        for paper in papers:
            comp = completeness_map.get(paper.student_name)
            eval_ = evaluation_map.get(paper.student_name)
            aigc = aigc_map.get(paper.student_name)
            plag = plagiarism_map.get(paper.student_name)

            lines.append(f"### {paper.student_name} — {paper.title}")
            lines.append("")
            lines.append(f"**基本信息**：{paper.word_count}字 ｜ {paper.paragraph_count}段落 ｜ 图片{paper.figure_count} / 表格{paper.table_count}")
            lines.append("")

            # 完整性
            if comp:
                lines.append(f"**完整性检测**：{comp.score:.0f}/100")
                if comp.missing_sections:
                    lines.append(f"- 缺失部分：{', '.join(comp.missing_sections)}")
                for w in comp.warnings:
                    lines.append(f"- ⚠ {w}")
                lines.append("")

            # AI评审
            if eval_ and eval_.success:
                lines.append(f"**AI评审总分**：{eval_.total_score}/100")
                lines.append("")
                if eval_.dimension_scores:
                    lines.append("| 维度 | 得分 |")
                    lines.append("|------|------|")
                    for dn, sc in eval_.dimension_scores.items():
                        lines.append(f"| {dn} | {sc} |")
                    lines.append("")
                lines.append(f"**评语**：{eval_.short_comment}")
                lines.append("")
                lines.append(f"**评分依据**：{eval_.evaluation_basis}")
                lines.append("")
            elif eval_:
                lines.append(f"**AI评审失败**：{eval_.error_message}")
                lines.append("")

            # AIGC
            if aigc:
                lines.append(f"**AIGC检测**：{aigc.overall_risk}（AI概率 {aigc.ai_probability:.0%}）")
                for seg in aigc.suspicious_segments[:3]:
                    lines.append(f"- [{seg['type']}] {seg['context']}")
                lines.append("")

            # 查重
            if plag and plag.suspicious_pairs:
                lines.append(f"**查重结果**：最高相似度 {plag.highest_similarity:.0%}（与 {plag.most_similar_student}）")
                lines.append("")
            else:
                lines.append("**查重结果**：✅ 无异常")
                lines.append("")

            lines.append("---")
            lines.append("")

        # 写入文件
        md_content = "\n".join(lines)
        output_path = self.output_dir / filename
        output_path.write_text(md_content, encoding="utf-8")
        print(f"[MD报告] 已保存至: {output_path}")
        return str(output_path)
