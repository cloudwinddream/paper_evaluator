"""
生成测试用论文 docx，含嵌入截图图片，用于验证全流程图片提取 + 分析。
"""
import io
from pathlib import Path
from PIL import Image, ImageDraw, ImageFont

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _make_screenshot(w: int, h: int, text: str, bg: tuple) -> bytes:
    """生成一张模拟程序截图的 PNG 字节"""
    img = Image.new("RGB", (w, h), color=bg)
    draw = ImageDraw.Draw(img)

    # 标题栏
    draw.rectangle([0, 0, w, 28], fill=(50, 50, 50))
    draw.rectangle([0, 28, w, 30], fill=(200, 200, 200))

    # 窗口按钮
    draw.ellipse([8, 8, 20, 20], fill=(255, 95, 87))
    draw.ellipse([24, 8, 36, 20], fill=(255, 189, 46))
    draw.ellipse([40, 8, 52, 20], fill=(39, 201, 63))

    # 文字居中
    try:
        font = ImageFont.truetype("/System/Library/Fonts/PingFang.ttc", 16)
    except OSError:
        font = ImageFont.load_default()

    bbox = draw.textbbox((0, 0), text, font=font)
    tx = (w - (bbox[2] - bbox[0])) // 2
    ty = (h - (bbox[3] - bbox[1])) // 2 + 10
    draw.text((tx, ty), text, fill=(30, 30, 30), font=font)

    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def _make_diagram(w: int, h: int) -> bytes:
    """生成一张模拟架构图的 PNG 字节"""
    img = Image.new("RGB", (w, h), color="white")
    draw = ImageDraw.Draw(img)

    boxes = [
        (50, 30, 200, 80, "数据采集", (173, 216, 230)),
        (280, 30, 430, 80, "数据存储", (144, 238, 144)),
        (160, 110, 310, 160, "数据处理", (255, 182, 193)),
        (50, 190, 200, 240, "数据分析", (255, 255, 150)),
        (280, 190, 430, 240, "可视化", (216, 191, 216)),
    ]

    try:
        font = ImageFont.truetype("/System/Library/Fonts/PingFang.ttc", 14)
    except OSError:
        font = ImageFont.load_default()

    for x1, y1, x2, y2, label, color in boxes:
        draw.rectangle([x1, y1, x2, y2], fill=color, outline="black", width=2)
        bbox = draw.textbbox((0, 0), label, font=font)
        tx = (x1 + x2) // 2 - (bbox[2] - bbox[0]) // 2
        ty = (y1 + y2) // 2 - (bbox[3] - bbox[1]) // 2
        draw.text((tx, ty), label, fill="black", font=font)

    for start, end in [((150, 80), (150, 110)), ((355, 80), (355, 110)),
                       ((125, 160), (125, 190)), ((300, 160), (300, 190))]:
        draw.line([start, end], fill="gray", width=2)

    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def create_test_docx(output_path: str):
    """生成一篇含截图的测试论文"""
    doc = Document()

    title = doc.add_heading("基于 Hadoop 生态的股票数据分析平台", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("学生：张三\n学号：2024001\n日期：2026-06-12")
    doc.add_paragraph("")

    # 第1章
    doc.add_heading("1. 项目概述", level=1)
    doc.add_paragraph(
        "本项目基于 Hadoop 生态构建了一个海量股票行情数据分析平台。"
        "系统采用 HDFS 进行数据存储，MapReduce 进行数据清洗和计算，"
        "Hive 进行数据仓库分析，Sqoop 进行数据导入导出。"
        "前端使用 ECharts 进行可视化展示。"
    )

    doc.add_heading("1.1 系统架构", level=2)
    doc.add_paragraph(
        "系统的整体架构如下图所示，包括数据采集层、数据存储层、"
        "数据处理层、数据分析层和可视化层五个模块。"
    )

    diagram_png = _make_diagram(480, 280)
    doc.add_picture(io.BytesIO(diagram_png), width=Inches(5.0))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("图1 系统架构图").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "从架构图可以看出，系统采用分层设计，各模块职责清晰，"
        "数据流从采集到可视化形成完整闭环。"
    )

    # 第2章
    doc.add_heading("2. 功能实现", level=1)

    doc.add_heading("2.1 数据采集模块", level=2)
    doc.add_paragraph(
        "数据采集模块负责从 Yahoo Finance 获取股票历史数据。"
        "使用 Python 编写爬虫脚本，每天定时采集数据并上传到 HDFS。"
    )

    screenshot1 = _make_screenshot(640, 360, "HDFS 文件系统界面 - /user/stock/data", (240, 240, 245))
    doc.add_picture(io.BytesIO(screenshot1), width=Inches(5.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("图2 HDFS 数据存储目录").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "上图展示了 HDFS 中存储的股票数据目录结构。"
    )

    doc.add_heading("2.2 数据清洗与处理", level=2)
    doc.add_paragraph(
        "使用 MapReduce 对原始数据进行清洗，包括去重、异常值处理、"
        "缺失值填充等操作。"
    )

    screenshot2 = _make_screenshot(640, 360, "MapReduce Job 运行状态 - Job_20260612_001", (245, 245, 220))
    doc.add_picture(io.BytesIO(screenshot2), width=Inches(5.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("图3 MapReduce 任务运行状态").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 第3章
    doc.add_heading("3. 实验结果", level=1)
    doc.add_paragraph(
        "实验使用 2010-2025 年的股票交易数据，共约 5 亿条记录。"
    )

    screenshot3 = _make_screenshot(640, 360, "Hive 查询结果 - SELECT * FROM stock_analysis LIMIT 10", (230, 240, 250))
    doc.add_picture(io.BytesIO(screenshot3), width=Inches(5.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("图4 Hive 查询结果示例").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 第4章
    doc.add_heading("4. 核心代码实现", level=1)
    doc.add_paragraph("以下是 MapReduce 的核心代码。")

    screenshot4 = _make_screenshot(640, 400, "MapReduce Java 代码 - StockMapper.java", (250, 250, 250))
    doc.add_picture(io.BytesIO(screenshot4), width=Inches(5.5))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("图5 MapReduce Mapper 代码").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("5. 总结与展望", level=1)
    doc.add_paragraph(
        "本文设计并实现了基于 Hadoop 生态的股票数据分析平台。"
        "未来可以引入 Spark 进行实时计算。"
    )

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"  ✓ 测试论文已生成: {output_path}")


if __name__ == "__main__":
    create_test_docx("papers/测试论文_张三_基于Hadoop的股票数据分析平台.docx")
