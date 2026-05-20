"""
创建测试用的示例论文
"""

from docx import Document
from docx.shared import Pt, Cm
from pathlib import Path


def create_sample_paper(filepath: str, student_name: str, title: str, content: dict):
    """创建示例论文"""
    doc = Document()

    # 标题
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.font.size = Pt(18)
    title_run.bold = True
    title_para.alignment = 1  # 居中

    # 摘要
    if "abstract" in content:
        doc.add_heading("摘要", level=2)
        doc.add_paragraph(content["abstract"])

    # 关键词
    if "keywords" in content:
        doc.add_heading("关键词", level=2)
        doc.add_paragraph(content["keywords"])

    # 引言
    if "introduction" in content:
        doc.add_heading("引言", level=2)
        doc.add_paragraph(content["introduction"])

    # 正文
    if "body" in content:
        doc.add_heading("正文", level=2)
        for section in content["body"]:
            doc.add_heading(section.get("title", ""), level=3)
            doc.add_paragraph(section.get("content", ""))

    # 结论
    if "conclusion" in content:
        doc.add_heading("结论", level=2)
        doc.add_paragraph(content["conclusion"])

    # 参考文献
    if "references" in content:
        doc.add_heading("参考文献", level=2)
        for ref in content["references"]:
            doc.add_paragraph(ref, style="List Number")

    doc.save(filepath)
    print(f"  创建: {filepath}")


def main():
    output_dir = Path("./test_papers")
    output_dir.mkdir(exist_ok=True)

    # 论文1：正常论文
    create_sample_paper(
        str(output_dir / "张三_人工智能论文.docx"),
        "张三",
        "人工智能在医疗领域的应用研究",
        {
            "abstract": "随着人工智能技术的快速发展，其在医疗领域的应用日益广泛。本文从人工智能的基本概念出发，分析了AI在医学影像诊断、药物研发、健康管理等方面的应用现状，并探讨了其面临的挑战和未来发展趋势。研究表明，人工智能在提高医疗效率、降低医疗成本方面具有巨大潜力。",
            "keywords": "人工智能；医疗；医学影像；药物研发",
            "introduction": "人工智能（Artificial Intelligence，AI）是计算机科学的一个重要分支，旨在研究和开发能够模拟、延伸和扩展人类智能的理论、方法和技术。近年来，随着深度学习、自然语言处理等技术的突破，人工智能在各个领域都取得了显著进展。医疗健康领域作为人工智能应用的重要场景，正经历着深刻的变革。",
            "body": [
                {
                    "title": "人工智能在医学影像诊断中的应用",
                    "content": "医学影像诊断是人工智能在医疗领域最成熟的应用之一。通过深度学习算法，AI可以快速、准确地分析X光片、CT、MRI等医学影像，辅助医生进行疾病诊断。例如，在肺部CT影像分析中，AI系统可以自动识别肺结节、肺炎等病变，诊断准确率已达到甚至超过资深放射科医生的水平。这不仅提高了诊断效率，也降低了漏诊和误诊的风险。"
                },
                {
                    "title": "人工智能在药物研发中的应用",
                    "content": "传统的药物研发过程耗时长、成本高，通常需要10-15年的时间和数十亿美元的投入。人工智能的引入正在改变这一局面。AI可以通过分析大量的分子结构数据，快速筛选出具有潜在药效的化合物，大大缩短了药物发现的时间。此外，AI还可以预测药物的副作用，优化临床试验设计，提高研发成功率。"
                },
                {
                    "title": "面临的挑战与未来展望",
                    "content": "尽管人工智能在医疗领域展现出巨大潜力，但其广泛应用仍面临诸多挑战。首先是数据隐私和安全问题，医疗数据涉及患者隐私，如何在保护隐私的同时充分利用数据是一个重要课题。其次是算法的可解释性，医生需要了解AI做出诊断的依据。最后是监管和伦理问题，需要建立相应的法律法规来规范AI在医疗中的应用。"
                }
            ],
            "conclusion": "综上所述，人工智能在医疗领域的应用前景广阔，正在从辅助诊断向精准医疗、个性化治疗等更深层次发展。未来，随着技术的不断成熟和政策的逐步完善，人工智能将为医疗健康事业带来更大的变革。我们应该积极拥抱这一技术变革，同时审慎应对其带来的挑战。",
            "references": [
                "[1] 李明, 王华. 人工智能在医学影像诊断中的应用进展[J]. 中国医学影像技术, 2023, 39(5): 789-793.",
                "[2] 张伟, 刘洋. 深度学习在药物研发中的应用[J]. 中国药学杂志, 2023, 58(3): 234-240.",
                "[3] 陈静, 赵强. 医疗人工智能的伦理与法律问题研究[J]. 医学与哲学, 2023, 44(2): 45-49.",
                "[4] Smith J, Brown A. AI in Healthcare: Past, Present and Future[J]. Nature Medicine, 2022, 28(6): 1120-1128.",
                "[5] 周磊, 吴敏. 人工智能辅助诊断系统的临床验证研究[J]. 中华医学杂志, 2023, 103(15): 1123-1128."
            ]
        }
    )

    # 论文2：缺少部分内容
    create_sample_paper(
        str(output_dir / "李四_人工智能论文.docx"),
        "李四",
        "人工智能技术发展概述",
        {
            "abstract": "本文简要介绍了人工智能的发展历程和主要技术方向。",
            "keywords": "人工智能；机器学习；深度学习",
            "introduction": "人工智能是当今科技领域最热门的话题之一。",
            "body": [
                {
                    "title": "发展历程",
                    "content": "人工智能的发展可以追溯到20世纪50年代。从早期的符号主义到现在的连接主义，AI经历了几次起伏。近年来，随着大数据和算力的提升，深度学习技术取得了突破性进展。"
                }
            ],
            "conclusion": "人工智能技术正在快速发展，未来将有更广泛的应用。",
            "references": [
                "[1] 王大明. 人工智能简史[M]. 北京: 科技出版社, 2022."
            ]
        }
    )

    # 论文3：AI风格明显的论文
    create_sample_paper(
        str(output_dir / "王五_人工智能论文.docx"),
        "王五",
        "人工智能的发展与应用研究",
        {
            "abstract": "首先，我们需要明确人工智能的基本概念。在当今社会，人工智能已经成为科技发展的重要方向。本文将从多个角度分析人工智能的发展现状和应用前景。不言而喻，人工智能将对人类社会产生深远影响。",
            "keywords": "人工智能；发展；应用；研究",
            "introduction": "众所周知，人工智能是计算机科学的一个重要分支。随着科技的不断发展，人工智能技术已经渗透到我们生活的方方面面。在本文中，我们将详细探讨人工智能的历史发展、核心技术以及未来趋势。",
            "body": [
                {
                    "title": "历史发展",
                    "content": "人工智能的发展可以追溯到20世纪50年代。首先，我们需要明确的是，人工智能的概念最早由达特茅斯会议提出。从那时起，人工智能经历了几次发展浪潮。值得注意的是，每一次技术突破都带来了新的应用场景。"
                },
                {
                    "title": "核心技术",
                    "content": "人工智能的核心技术包括机器学习、深度学习和自然语言处理等。一方面，这些技术为AI的发展提供了基础；另一方面，它们也在不断演进。不可否认的是，深度学习是当前最热门的技术方向。"
                },
                {
                    "title": "应用领域",
                    "content": "人工智能在各个领域都有广泛应用。在医疗领域，AI可以辅助诊断；在交通领域，AI可以优化路线规划；在教育领域，AI可以实现个性化学习。由此可见，人工智能的应用前景非常广阔。"
                }
            ],
            "conclusion": "综上所述，我们可以得出结论：人工智能是一项具有巨大潜力的技术。未来，人工智能将会在更多领域发挥重要作用。我们应该积极拥抱这一技术变革，同时也要关注其带来的挑战。相信在不久的将来，人工智能将会为人类社会带来更多福祉。",
            "references": [
                "[1] 人工智能概论. 科技出版社.",
                "[2] 深度学习原理与实践. 电子工业出版社.",
                "[3] 机器学习算法详解. 清华大学出版社.",
                "[4] 自然语言处理导论. 人民邮电出版社.",
                "[5] 人工智能应用案例集. 机械工业出版社."
            ]
        }
    )

    # 论文4：与论文1有部分内容重复（模拟抄袭）
    create_sample_paper(
        str(output_dir / "赵六_人工智能论文.docx"),
        "赵六",
        "AI在医疗健康领域的应用",
        {
            "abstract": "随着人工智能技术的快速发展，其在医疗领域的应用日益广泛。本文分析了AI在医学影像诊断、药物研发等方面的应用现状。研究表明，人工智能在提高医疗效率方面具有巨大潜力。",
            "keywords": "AI；医疗健康；医学影像",
            "introduction": "人工智能（Artificial Intelligence，AI）是计算机科学的一个重要分支。医疗健康领域作为人工智能应用的重要场景，正经历着深刻的变革。近年来，随着深度学习技术的突破，AI在医疗领域取得了显著进展。",
            "body": [
                {
                    "title": "医学影像诊断",
                    "content": "医学影像诊断是人工智能在医疗领域最成熟的应用之一。通过深度学习算法，AI可以快速、准确地分析X光片、CT、MRI等医学影像，辅助医生进行疾病诊断。在肺部CT影像分析中，AI系统可以自动识别肺结节、肺炎等病变，诊断准确率已达到甚至超过资深放射科医生的水平。"
                },
                {
                    "title": "药物研发",
                    "content": "传统的药物研发过程耗时长、成本高。人工智能的引入正在改变这一局面。AI可以通过分析大量的分子结构数据，快速筛选出具有潜在药效的化合物，大大缩短了药物发现的时间。此外，AI还可以预测药物的副作用，优化临床试验设计。"
                }
            ],
            "conclusion": "人工智能在医疗领域的应用前景广阔。未来，随着技术的不断成熟，人工智能将为医疗健康事业带来更大的变革。",
            "references": [
                "[1] 李明, 王华. 人工智能在医学影像诊断中的应用进展[J]. 中国医学影像技术, 2023.",
                "[2] 张伟. 深度学习在药物研发中的应用[J]. 中国药学杂志, 2023.",
                "[3] 陈静. 医疗人工智能的伦理问题研究[J]. 医学与哲学, 2023.",
                "[4] Smith J. AI in Healthcare[J]. Nature Medicine, 2022.",
                "[5] 周磊. 人工智能辅助诊断系统研究[J]. 中华医学杂志, 2023."
            ]
        }
    )

    print(f"\n已创建4篇测试论文到 {output_dir} 目录")


if __name__ == "__main__":
    main()
